from tkinter import *
from tkinter import filedialog
from pydub import AudioSegment, silence
from pydub.silence import detect_silence, split_on_silence
from docx import Document
import speech_recognition as sr
import os

import requests
import json
 
def get_fileName():
    global fileName
    fileName = filedialog.askopenfilename(parent=file_window, defaultextension='.mp3')
    file_window.quit()
    file_window.destroy()
    return fileName

def translate(text, source='en', target='ko'):
    CLIENT_ID, CLIENT_SECRET = 'bIkqYIrqHLRMe7NC7Fmi', 'ZPRAbnzXFe'
    url = 'https://openapi.naver.com/v1/papago/n2mt'
    headers = {
        'Content-Type': 'application/json',
        'X-Naver-Client-Id': CLIENT_ID,
        'X-Naver-Client-Secret': CLIENT_SECRET
    }
    data = {'source': 'en', 'target': 'ko', 'text': text}
    response = requests.post(url, json.dumps(data), headers=headers)
    return response.json()['message']['result']['translatedText']

# File Selection Winow
file_window = Tk()
file_window.geometry("300x50")
file_window.resizable(width=False, height=False)
file_window.title("Audio File")
Label(file_window, text="File").grid(row=0, padx=10, pady=10)
Button(file_window, text="파일 선택", command=get_fileName).grid(row=0, column=1, padx=10, pady=10)
file_window.mainloop()

# creating new document
doc = Document()
doc.add_heading('Listening Script', level=0)

sound = AudioSegment.from_mp3(fileName)

# spliting audio files
audio_chunks = split_on_silence(sound, min_silence_len=2000, silence_thresh=-100 )

for i, chunk in enumerate(audio_chunks, start=1):
   section_count = i
   output_file = "section{0}.wav".format(i)
   print("Exporting file", output_file)
   chunk.export(output_file, format="wav")

for i in range(1, section_count + 1):
   section_file_name = "section{0}.wav".format(i)

   # Speech Recognition
   r = sr.Recognizer()
   section_file = sr.AudioFile(section_file_name)
   with section_file as source:
      audio = r.record(source)

   result = r.recognize_google(audio)
   translted_result = translate(result)

   sectionNUM = doc.add_paragraph()
   sectionNUM.add_run('Section ' + str(i)).bold = True

   doc.add_paragraph(result)
   doc.add_paragraph(translted_result + '\n')

doc.save('Listening-Script.docx')